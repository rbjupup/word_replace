
import os
import win32com.client
import pythoncom
import psutil
import json
from docx import Document
USE_DIR = "D:/DATA/Tool/"
def replace_text_in_docx(input_file, output_file, replacements):
    # 读取文档
    doc = Document(input_file)

    # 替换段落中的文本
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
    # 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 获取单元格中所有运行的文本
                for old_text, new_text in replacements.items():
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
    # 保存新的文档
    doc.save(output_file)
    





def convert_docx_to_pdf(docx_file, pdf_file):
    pythoncom.CoInitialize()  # 初始化 COM 库（多线程环境必须）
    word = None
    doc = None
    try:
        # 转换核心逻辑
        word = win32com.client.DispatchEx('Word.Application')  # 使用独立进程
        word.Visible = False  # 不可见模式提高稳定性
        docx_abs = os.path.abspath(docx_file)
        pdf_abs = os.path.abspath(pdf_file)
        
        doc = word.Documents.Open(docx_abs)
        doc.SaveAs(pdf_abs, FileFormat=17)  # 17 对应 PDF 格式
        
    except Exception as e:
        print(f"转换失败: {str(e)}")
        raise
    finally:
        # 确保资源释放
        try:
            if doc:
                doc.Close(SaveChanges=0)  # 0=wdDoNotSaveChanges
        except AttributeError:
            pass  # 忽略对象已失效的情况
        try:
            if word:
                pass#word.Quit()  # 先尝试正常退出
        except AttributeError:
            pass  # 忽略对象已失效的情况
        # 强制终止残留进程
        kill_word_process()
        pythoncom.CoUninitialize()  # 释放 COM 资源
        
def kill_word_process():
    # 终结所有残留的 WINWORD.EXE 进程
    for proc in psutil.process_iter(['pid', 'name']):
        if proc.info['name'] == 'WINWORD.EXE':
            try:
                proc.kill()
            except:
                pass
class App:
    def __init__(self):
        self.replacements = []
        self.DATA_FILE = USE_DIR+'/replacements.json'
        self.load_data()

    def load_data(self):
        if os.path.exists(self.DATA_FILE):
            with open(self.DATA_FILE, 'r', encoding='utf-8') as file:
                self.replacements = json.load(file)
        
    def Run(self):
        input_file = USE_DIR+r'/模板.docx'  # 输入文档路径
        output_file = USE_DIR+r'/Contract欧代协议-.docx'  # 输出文档路径
        pdf_file = USE_DIR+r'/Contract欧代协议-.pdf'  # 输出文档路径

        matching_items = [item for item in self.replacements.items() if item[0] == 'Replace_CompanyName_ZHCN']
        if matching_items:
            output_file = f'D:/DATA/Result/Contract欧代协议-{matching_items[0][1]}.docx'  # 输出文档路径
            pdf_file = f'D:/DATA/Result/Contract欧代协议-{matching_items[0][1]}.pdf'  # 输出文档路径
        replacements = {}
        for a, b in self.replacements.items():
            replacements[a] = b
        replace_text_in_docx(input_file, output_file, replacements)
        convert_docx_to_pdf(output_file, pdf_file)
if __name__ == "__main__":
    App().Run()