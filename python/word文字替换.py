import tkinter as tk
from tkinter import simpledialog, messagebox, Listbox
import json
import os
from docx import Document
import os
import win32com.client
import pandas as pd
from tkinter import filedialog

def replace_text_in_docx(input_file, output_file, replacements):
    doc = Document(input_file)

    # 替换段落中的文本
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, new_text)

    # 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for old_text, new_text in replacements.items():
                    for paragraph in cell.paragraphs:
                        if old_text in paragraph.text:
                            paragraph.text = paragraph.text.replace(old_text, new_text)

    # 替换页眉中的文本
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)

    # 替换页脚中的文本
    for section in doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            for old_text, new_text in replacements.items():
                if old_text in paragraph.text:
                    paragraph.text = paragraph.text.replace(old_text, new_text)

    # 保存新的文档
    doc.save(output_file)




def convert_docx_to_pdf(docx_file, pdf_file):
    # 使用 win32com 将 Word 转换为 PDF
    word = win32com.client.Dispatch('Word.Application')
    doc = word.Documents.Open(os.path.abspath(docx_file))
    doc.SaveAs(os.path.abspath(pdf_file), FileFormat=17)  # 17表示 PDF 格式
    doc.Close()
    word.Quit()

# 数据文件
DATA_FILE = 'replacements.json'

class App:
    def __init__(self, root):
        self.root = root
        self.root.title("文本替换工具")

        self.replacements = []

        self.listbox = Listbox(root, width=50)
        self.listbox.pack(pady=10)

        self.entry_a = tk.Entry(root, width=20)
        self.entry_a.pack(pady=5)
        self.entry_a.insert(0, "输入 A被替换字段")

        self.entry_b = tk.Entry(root, width=20)
        self.entry_b.pack(pady=5)
        self.entry_b.insert(0, "输入 B替换字段")

        self.add_button = tk.Button(root, text="添加替换", command=self.add_replacement)
        self.add_button.pack(pady=5)

        self.delete_button = tk.Button(root, text="删除选中", command=self.delete_replacement)
        self.delete_button.pack(pady=5)

        self.import_button = tk.Button(root, text="从excel导入", command=self.start_import)
        self.import_button.pack(pady=5)

        self.replace_button = tk.Button(root, text="开始替换", command=self.start_replacement)
        self.replace_button.pack(pady=5)

        self.load_data()
        self.listbox.bind('<Double-1>', self.edit_replacement)

    def add_replacement(self):
        a = self.entry_a.get().strip()
        b = self.entry_b.get().strip()
        if a and b:
            self.replacements.append((a, b))
            self.listbox.insert(tk.END, f"{a} -> {b}")
            self.save_data()
            self.entry_a.delete(0, tk.END)
            self.entry_b.delete(0, tk.END)
        else:
            messagebox.showwarning("警告", "A 和 B 不能为空")

    def delete_replacement(self):
        selected_index = self.listbox.curselection()
        if selected_index:
            index = selected_index[0]
            self.replacements.pop(index)  # 从列表中删除
            self.listbox.delete(index)     # 从界面中删除
            self.save_data()
        else:
            messagebox.showwarning("警告", "请先选择要删除的项")

    def start_replacement(self):
        input_file = r'模板.docx'  # 输入文档路径
        output_file = r'Contract欧代协议-.docx'  # 输出文档路径
        pdf_file = r'Contract欧代协议-.pdf'  # 输出文档路径

        matching_items = [item for item in self.replacements if item[0] == 'dataC3']
        if matching_items:
            output_file = f'Contract欧代协议-{matching_items[0][1]}.docx'  # 输出文档路径
            pdf_file = f'Contract欧代协议-{matching_items[0][1]}.pdf'  # 输出文档路径
        replacements = {}
        for a, b in self.replacements:
            replacements[a] = b
        replace_text_in_docx(input_file, output_file, replacements)
        convert_docx_to_pdf(output_file, pdf_file)
        messagebox.showinfo("处理完成","替换完成,已将{}替换为{}".format(input_file, output_file))

    def start_import(self):
        # 弹出文件选择对话框，选择xlsx文件
        file_path = filedialog.askopenfilename(
            title="选择文件",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")]
        )
        if file_path:
            # 打印选择的文件名称
            print(f"选择的文件: {file_path}")
            # 读取Excel文件
            df = pd.read_excel(file_path)
            dataC3 = df.iloc[1, 2]
            dataC4 = df.iloc[2, 2]
            dataC5 = df.iloc[3, 2]
            dataC7 = df.iloc[5, 2]
            dataC8 = df.iloc[6, 2]
            dataC9 = df.iloc[7, 2]
            dataC10 = df.iloc[8, 2]
            self.add_or_update('dataC3',dataC3)
            self.add_or_update('dataC4',dataC4)
            self.add_or_update('dataC5',dataC5)
            self.add_or_update('dataC7',dataC7)
            self.add_or_update('dataC8',dataC8)
            self.add_or_update('dataC9',dataC9)
            self.add_or_update('dataC10',dataC10)

    def add_or_update(self,value,value2):
        for i, item in enumerate(self.replacements):
            if item[0] == value:  # 假设我们检查第一列的值
                # 更新存在的项
                self.replacements[i] = (value, value2)  # 示例更新
                self.update_listbox()
                #messagebox.showinfo("Update", f"Updated: {value}")
                return

        # 如果不存在则添加
        self.replacements.append((value, value2))
        self.update_listbox()
        #messagebox.showinfo("Add", f"Added: {value}")

    def update_listbox(self):
        self.listbox.delete(0, tk.END)  # 清空Listbox
        # for item in self.replacements:
        #     self.listbox.insert(tk.END, item)  # 插入更新后的数据
        for a, b in self.replacements:
            self.listbox.insert(tk.END, f"{a} -> {b}")
        self.save_data()

    def edit_replacement(self, event):
        selected_index = self.listbox.curselection()
        if selected_index:
            index = selected_index[0]
            a, b = self.replacements[index]
            new_b = simpledialog.askstring("修改 B", f"当前 B 值为 {b}, 输入新的 B 值:")
            if new_b is not None:
                self.replacements[index] = (a, new_b)
                self.listbox.delete(index)
                self.listbox.insert(index, f"{a} -> {new_b}")
                self.save_data()

    def load_data(self):
        if os.path.exists(DATA_FILE):
            with open(DATA_FILE, 'r', encoding='utf-8') as file:
                self.replacements = json.load(file)
                for a, b in self.replacements:
                    self.listbox.insert(tk.END, f"{a} -> {b}")

    def save_data(self):
        with open(DATA_FILE, 'w', encoding='utf-8') as file:
            json.dump(self.replacements, file)

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()
