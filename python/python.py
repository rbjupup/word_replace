from docx import Document
import os
import win32com.client

def replace_text_in_docx(input_file, output_file, replacements):
    # 读取文档
    doc = Document(input_file)

    # 替换段落中的文本
    for paragraph in doc.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)
# 替换表格中的文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 获取单元格中所有运行的文本
                for old_text, new_text in replacements.items():
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
    # # 替换表格中的文本
    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             print(cell.text)
    #             for old_text, new_text in replacements.items():
    #                 if old_text in cell.text:
    #                     for run in cell.paragraphs[0].runs:
    #                         if old_text in run.text:
    #                             run.text = run.text.replace(old_text, new_text)

    # 保存新的文档
    doc.save(output_file)

def convert_docx_to_pdf(docx_file, pdf_file):
    # 使用 win32com 将 Word 转换为 PDF
    word = win32com.client.Dispatch('Word.Application')
    doc = word.Documents.Open(os.path.abspath(docx_file))
    doc.SaveAs(os.path.abspath(pdf_file), FileFormat=17)  # 17表示 PDF 格式
    doc.Close()
    word.Quit()

# 使用示例
input_file = r'模板.docx'  # 输入文档路径
output_file = r'结果.docx'  # 输出文档路径
pdf_file = r'结果.pdf'  # 输出文档路径

replacements = {
    '姓名': 'Yangkunzhang',
    '地址': 'shenzhen jingchuang gangzhilong',
    '电话': '188xxxxxxxx',
    '邮箱': 'yangkunzhang@ptcaoi.com',
    '法人名字': 'Yangkunzhang'
}

replace_text_in_docx(input_file, output_file, replacements)
convert_docx_to_pdf(output_file, pdf_file)
